from fnmatch import translate
from pylatex import Document, Section, Subsection, Command
from pylatex.base_classes import Environment
from docx import Document as docx_Document
from translate import Zh2En

from pylatex.utils import italic, NoEscape
def fill_document(doc):
    """Add a section, a subsection and some text to the document.

    :param doc: the document
    :type doc: :class:`pylatex.document.Document` instance
    """
    with doc.create(Section('A section')):
        doc.append('Some regular text and some ')
        doc.append(italic('italic text. '))

        with doc.create(Subsection('A subsection')):
            doc.append('Also some crazy characters: $&#{}')

class ABSTRACT(Environment):
    """A class to wrap LaTeX's alltt environment."""

    # packages = [Package('alltt')]
    escape = False
    content_separator = "\n"
if __name__ == '__main__':
    # Basic document
    doc = Document(default_filepath='paper',documentclass='IEEEtran',document_options='journal',lmodern=False,textcomp=False)
    doc.preamble.append(Command('title', 'Translate to English using the YNMT natural language translation model'))
    doc.preamble.append(Command('author', '***, ***'))
    
    doc.append(NoEscape(r'\maketitle'))
    with doc.create(ABSTRACT()):
         
        abstract = docx_Document(r"./abstract.docx")
        for p in abstract.paragraphs:
            abstrach_zh = p.text
            abstrach_en = Zh2En(abstrach_zh)

        
            doc.append(abstrach_en)
    with doc.create(Section('Introduction')):
        introduction = docx_Document(r"./introduction.docx")
        k=1
        for p in introduction.paragraphs:
            introduction_zh = p.text
            introduction_en = Zh2En(introduction_zh)
            if k==1:
                introduction_en[0] = "\IEEEPARstart{"+introduction_en[0]+"}" 
            doc.append(abstrach_en)
    
    
    
  

    doc.generate_pdf(clean_tex=False)
    doc.generate_tex()

    # Document with `\maketitle` command activated
    # doc = Document()

    # doc.preamble.append(Command('title', 'Awesome Title'))
    # doc.preamble.append(Command('author', 'Anonymous author'))
    # doc.preamble.append(Command('date', NoEscape(r'\today')))
    # doc.append(NoEscape(r'\maketitle'))

    # fill_document(doc)

    # doc.generate_pdf('basic_maketitle', clean_tex=False)

    # # Add stuff to the document
    # with doc.create(Section('A second section')):
    #     doc.append('Some text.')

    # doc.generate_pdf('basic_maketitle2', clean_tex=False)
    # tex = doc.dumps()  # The document as string in LaTeX syntax